import uuid
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from app.config import settings


def generate_professional_docx(
    title: str,
    subject: str,
    executive_summary: str,
    findings: list[dict],
    calculations: list[dict],
    recommendations: list[str],
    references: list[str],
    document_id: str | None = None
) -> str:
    """Creates a high-quality, formatted Word document (.docx) conforming to industrial approval note standards."""
    doc_id = document_id or f"MRPL-APPR-{datetime.utcnow().strftime('%Y%m%d')}-{uuid.uuid4().hex[:4].upper()}"
    file_name = f"Approval_Note_{doc_id}.docx"
    file_path = settings.ARTIFACTS_DIR / file_name

    doc = Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("MANGALORE REFINERY AND PETROCHEMICALS LIMITED")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 44, 89)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("TECHNICAL INSPECTION & COMPLIANCE APPROVAL NOTE")
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(12)
    sub_run.font.bold = True

    # Metadata Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Document ID:", doc_id),
        ("Subject / Asset:", subject),
        ("Evaluation Date:", datetime.utcnow().strftime("%d-%b-%Y")),
        ("Status:", "AI Draft — Human Technical Verification Required")
    ]
    for row_idx, (k, v) in enumerate(meta_data):
        cell_k = meta_table.cell(row_idx, 0)
        cell_v = meta_table.cell(row_idx, 1)
        cell_k.paragraphs[0].add_run(k).font.bold = True
        cell_v.paragraphs[0].add_run(v)

    doc.add_paragraph()  # spacing

    # 1. Executive Summary
    h1 = doc.add_heading("1. Executive Summary", level=1)
    p_summary = doc.add_paragraph(executive_summary)
    p_summary.paragraph_format.line_spacing = 1.15

    # 2. Key Inspection Findings
    doc.add_heading("2. Technical Findings & Condition Assessment", level=1)
    if findings:
        findings_table = doc.add_table(rows=1, cols=4)
        findings_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = findings_table.rows[0].cells
        hdr_cells[0].text = "Asset / Tag"
        hdr_cells[1].text = "Observed Condition"
        hdr_cells[2].text = "Measured Value"
        hdr_cells[3].text = "Severity"
        for c in hdr_cells:
            c.paragraphs[0].runs[0].font.bold = True
            # Set background color for header
            shading = parse_xml(r'<w:shd {} w:fill="1E3A8A"/>'.format(nsdecls('w')))
            c._tc.get_or_add_tcPr().append(shading)
            c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        for f in findings:
            row_cells = findings_table.add_row().cells
            row_cells[0].text = str(f.get("tag", "N/A"))
            row_cells[1].text = str(f.get("condition", "N/A"))
            row_cells[2].text = str(f.get("measured", "N/A"))
            row_cells[3].text = str(f.get("severity", "Normal"))

    # 3. Deterministic Engineering Calculations
    doc.add_heading("3. Deterministic Engineering Calculations", level=1)
    if calculations:
        for c in calculations:
            calc_p = doc.add_paragraph()
            calc_p.add_run(f"• Parameter: {c.get('parameter', 'N/A')}\n").font.bold = True
            calc_p.add_run(f"  Formula: {c.get('formula', 'N/A')}\n")
            calc_p.add_run(f"  Result: {c.get('result', 'N/A')}\n")
            calc_p.add_run(f"  Standard Limit: {c.get('limit', 'N/A')}\n")

    # 4. Actionable Recommendations
    doc.add_heading("4. Recommended Action Plan", level=1)
    for rec in recommendations:
        doc.add_paragraph(f"• {rec}", style='List Bullet')

    # 5. Applicable References & SOP Citations
    doc.add_heading("5. Internal References & Standard Citations", level=1)
    for ref in references:
        doc.add_paragraph(f"• {ref}", style='List Bullet')

    # 6. Formal Sign-Off Block
    doc.add_heading("6. Human Approval & Sign-Off", level=1)
    sign_table = doc.add_table(rows=2, cols=2)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign_table.cell(0, 0).text = "Prepared By (AI Agent Evaluator):\n\nSovereignAI Workbench v1.0\nDate: " + datetime.utcnow().strftime("%d-%b-%Y")
    sign_table.cell(0, 1).text = "Reviewed & Approved By (Lead Engineer):\n\nSignature: _______________________\nName:\nDate:"
    sign_table.cell(1, 0).text = "Operations In-Charge:\n\nSignature: _______________________"
    sign_table.cell(1, 1).text = "Safety & Inspection Department:\n\nSignature: _______________________"

    # Safety Notice Footer
    footer = doc.sections[0].footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = footer_p.add_run("CONFIDENTIAL & SOVEREIGN INDUSTRIAL DOCUMENT — GENERATED ON-PREMISE WITH LOCAL OPEN-WEIGHT AI — HUMAN APPROVAL MANDATORY")
    f_run.font.size = Pt(8)
    f_run.font.color.rgb = RGBColor(120, 120, 120)

    doc.save(str(file_path))
    return str(file_path)
